import reports
from errors import *
from docx import Document
from docx.shared import Cm
from uuid import uuid4
import utils
import avito
import datetime
import calendar
import utils


monthes = ['', 'Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь',
           'Декабрь']


def generate_report_by_all_cities(year, month):
    days_range = calendar.monthrange(year, month)
    date_from = datetime.datetime(year, month, 1)
    date_to = datetime.datetime(year, month, days_range[1])
    data = reports.api.get_avito_report(date_from, date_to)

    all_views = 0
    all_contacts = 0
    all_favorites = 0
    all_wmembers_list = []
    wmembers_list = []

    for ad in data:
        for stat in ad['stats']:
            all_views += stat['uniqViews']
            all_contacts += stat['uniqContacts']
            all_favorites += stat['uniqFavorites']
        wmembers_10 = list(filter(lambda wmember: wmember.status == 10, ad['wmembers']))
        wmembers_list.extend(wmembers_10)
        all_wmembers_list.extend(ad['wmembers'])
        ad['wmembers_10'] = wmembers_10

    data.sort(key=lambda ad: len(ad['wmembers_10']), reverse=True)

    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    now = datetime.datetime.now()

    document.add_heading('Отчет по всем актуальным объявлениям за {} {}'.format(monthes[month], year), level=0)
    document.add_heading('Дата формирования отчёта: {}'.format(now.strftime('%d.%m.%Y %H:%M')), level=1)
    document.add_heading('Общая информация', level=1)
    document.add_paragraph(f'За {monthes[month]} {year} работало {len(data)} {utils.get_norm_word("объявление", len(data))} по такому же количеству городов.')
    document.add_paragraph(f'По итогу, было получено:')
    document.add_paragraph(
        f'Уникальных просмотров - {all_views}', style='List Bullet'
    )
    document.add_paragraph(
        f'Предоставили свои контакты - {all_contacts} {utils.get_norm_word("человек", all_contacts)}',
        style='List Bullet'
    )
    document.add_paragraph(
        f'Добавили объявление в избранное - {all_favorites} {utils.get_norm_word("человек", all_favorites)}',
        style='List Bullet'
    )
    document.add_paragraph(
        f'Вступилили в группу - {len(wmembers_list)} {utils.get_norm_word("человек", (len(wmembers_list)))}', style='List Bullet'
    )

    document.add_heading('Информация по объявлениям', level=1)
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Объявление'
    hdr_cells[1].text = 'Опубликовано дней'
    hdr_cells[2].text = 'Просмотры'
    hdr_cells[3].text = 'Контакты'
    hdr_cells[4].text = 'Избранное'
    hdr_cells[5].text = 'Вступили в группу'

    for ad in data:
        all_views = 0
        all_contacts = 0
        all_favorites = 0
        for stat in ad['stats']:
            all_views += stat['uniqViews']
            all_contacts += stat['uniqContacts']
            all_favorites += stat['uniqFavorites']
        row_cells = table.add_row().cells
        row_cells[0].text = '{} {}'.format(ad['city'].name, ad['title'])
        row_cells[1].text = str(len(ad['stats']))
        row_cells[2].text = str(all_views)
        row_cells[3].text = str(all_contacts)
        row_cells[4].text = str(all_favorites)
        row_cells[5].text = str(len(ad['wmembers_10']))

    document.add_heading('Информация по статусам соискателей', level=1)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Статус'
    hdr_cells[1].text = 'Количество'

    statuses = {
        '[Анкетирование] Ожидаем имя': 0,
        '[Анкетирование] Ожидаем возраст': 1,
        '[Анкетирование] Ожидаем город': 2,
        '[Анкетирование] Подтверждение данных': 3,
        '[Анкетирование] Установка Telegram': 4,
        '[Добавление] Ожидает добавление города': 15,
        '[Добавление] Ссылка отправлена. Переход не совершён': 6,
        'Состоит в группе': 10,
        'Покинул группу': 11,
    }

    for status in statuses:
        wmembers_count = len(list(filter(lambda wmember: wmember.status == statuses[status], all_wmembers_list)))
        row_cells = table.add_row().cells
        row_cells[0].text = status
        row_cells[1].text = str(wmembers_count)

    document.add_heading('Информация по источникам соискателей', level=1)

    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Источник'
    hdr_cells[1].text = 'Количество всего'
    hdr_cells[2].text = 'Количество в группе'

    statuses = {
        'WhatsApp': 2,
        'Avito': 1,
    }

    for status in statuses:
        wmembers_count = len(list(filter(lambda wmember: wmember.source_id == statuses[status], all_wmembers_list)))
        wmembers_count_need = len(
            list(filter(lambda wmember: wmember.source_id == statuses[status] and wmember.status == 10, wmembers_list)))
        row_cells = table.add_row().cells
        row_cells[0].text = status
        row_cells[1].text = str(wmembers_count)
        row_cells[2].text = str(wmembers_count_need)

    path = '{}/reports/docs/{}_Все города_{}_{}.docx'.format(utils.get_script_dir(), year, monthes[month], uuid4())
    document.save(path)
    return path


def generateReportByCity(year, month, ad_id):
    city = reports.api.get_city_by_avito_id(ad_id)
    if city is None:
        raise IncorrectDataValue('Не удалось определить город по объявлению')

    days_range = calendar.monthrange(year, month)

    date_from = datetime.datetime(year, month, 1)
    date_to = datetime.datetime(year, month, days_range[1])
    full_data = reports.api.get_avito_report_by_ad(ad_id, date_from, date_to, load_members=True)
    data = full_data['data']
    wmembers = full_data['wmembers']
    work_days = data['result']['items'][0]['stats']
    all_views = 0
    all_contacts = 0
    all_favorites = 0
    all_wmembers_list = list(filter(lambda wmember: wmember.status == 10, wmembers))
    all_wmembers = len(all_wmembers_list)

    for day in work_days:
        all_views += day['uniqViews']
        all_contacts += day['uniqContacts']
        all_favorites += day['uniqFavorites']

    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    now = datetime.datetime.now()

    document.add_heading('Отчет по объявлению #{} за {} {} | {}'.format(ad_id, monthes[month], year, city.name), level=0)
    document.add_heading('Дата формирования отчёта: {}'.format(now.strftime('%d.%m.%Y %H:%M')), level=1)

    document.add_heading('Статистка по объявлению', level=1)

    document.add_paragraph(f'За {monthes[month]} {year}г. объявление было активно {len(work_days)} {utils.get_norm_word("день", len(work_days))}.')
    document.add_paragraph(f'За {len(work_days)} {utils.get_norm_word("день", len(work_days))} работы было получено:')
    document.add_paragraph(
        f'Уникальных просмотров - {all_views}', style='List Bullet'
    )
    document.add_paragraph(
        f'Предоставили свои контакты - {all_contacts} {utils.get_norm_word("человек", all_contacts)}', style='List Bullet'
    )
    document.add_paragraph(
        f'Добавили объявление в избранное - {all_favorites} {utils.get_norm_word("человек", all_favorites)}', style='List Bullet'
    )
    document.add_paragraph(
        f'Вступилили в группу - {all_wmembers} {utils.get_norm_word("человек", all_wmembers)}', style='List Bullet'
    )

    document.add_heading('Информация по дням публикаций', level=1)

    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дата'
    hdr_cells[1].text = 'Просмотры'
    hdr_cells[2].text = 'Контакты'
    hdr_cells[3].text = 'Избранное'
    hdr_cells[4].text = 'Вступили в группу'

    for day in work_days:
        day_date = datetime.datetime.strptime(day['date'], '%Y-%m-%d')
        wmembers_list = list(filter(lambda wmember: utils.compareDate(wmember.create_date, day_date), all_wmembers_list))

        row_cells = table.add_row().cells
        row_cells[0].text = day_date.strftime('%d.%m.%Y')
        row_cells[1].text = str(day['uniqViews'])
        row_cells[2].text = str(day['uniqContacts'])
        row_cells[3].text = str(day['uniqFavorites'])
        row_cells[4].text = str(len(wmembers_list))

    document.add_heading('Информация по статусам соискателей', level=1)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Статус'
    hdr_cells[1].text = 'Количество'

    statuses = {
        '[Анкетирование] Ожидаем имя': 0,
        '[Анкетирование] Ожидаем возраст': 1,
        '[Анкетирование] Ожидаем город': 2,
        '[Анкетирование] Подтверждение данных': 3,
        '[Анкетирование] Установка Telegram': 4,
        '[Добавление] Ожидает добавление города': 15,
        '[Добавление] Ссылка отправлена. Переход не совершён': 6,
        'Состоит в группе': 10,
        'Покинул группу': 11,
    }

    for status in statuses:
        wmembers_count = len(list(filter(lambda wmember: wmember.status == statuses[status], wmembers)))
        row_cells = table.add_row().cells
        row_cells[0].text = status
        row_cells[1].text = str(wmembers_count)

    document.add_heading('Информация по источникам соискателей', level=1)

    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Источник'
    hdr_cells[1].text = 'Количество всего'
    hdr_cells[2].text = 'Количество в группе'

    statuses = {
        'WhatsApp': 2,
        'Avito': 1,
    }

    for status in statuses:
        wmembers_count = len(list(filter(lambda wmember: wmember.source == statuses[status], wmembers)))
        wmembers_count_need = len(list(filter(lambda wmember: wmember.source == statuses[status] and wmember.status == 10, wmembers)))
        row_cells = table.add_row().cells
        row_cells[0].text = status
        row_cells[1].text = str(wmembers_count)
        row_cells[2].text = str(wmembers_count_need)


    path = '{}/reports/docs/{}_{}_{}_{}.docx'.format(utils.get_script_dir(), city.name, year, monthes[month], uuid4())
    document.save(path)
    return path